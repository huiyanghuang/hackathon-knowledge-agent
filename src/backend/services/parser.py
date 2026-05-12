import re
import uuid
from pathlib import Path
from typing import Optional

from pypdf import PdfReader

from models.schemas import Chapter, Textbook


# 只匹配一级章标题，不匹配小节（如 1.1、第一节）
CHAPTER_PATTERNS = [
    re.compile(r"^第\s*[一二三四五六七八九十百\d]+\s*章\s*\S"),   # 第X章 ...
    re.compile(r"^Chapter\s+\d+\b", re.IGNORECASE),                # Chapter N
]

HEADER_FOOTER_MAX_LEN = 80

# 用于合并过小章节：少于此字数的章节合并到下一章
MIN_CHAPTER_CHARS = 500


def _is_chapter_title(text: str) -> bool:
    text = text.strip()
    # 长度限制：章标题一般不超过30字
    if not text or len(text) > 30:
        return False
    return any(p.match(text) for p in CHAPTER_PATTERNS)


def _is_header_footer(text: str) -> bool:
    text = text.strip()
    return len(text) < HEADER_FOOTER_MAX_LEN and any([
        re.match(r"^\d+$", text),
        re.match(r"^第\s*\d+\s*页", text),
        re.match(r"^Page\s+\d+", text, re.IGNORECASE),
    ])


PAGE_GROUP = 25  # 每组页数，约等于医学教材一个大章


def parse_pdf(file_path: str, textbook_id: Optional[str] = None) -> Textbook:
    """
    Page-group strategy: group every PAGE_GROUP pages into one chapter.
    More robust than regex-based chapter detection for Chinese medical PDFs
    where pypdf layout extraction is unreliable.
    Detected chapter titles are used as labels when found; otherwise use page range.
    """
    path = Path(file_path)
    tid = textbook_id or str(uuid.uuid4())[:8]
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)

    # 1. Extract text per page, detect chapter title on each page
    page_texts: list[str] = []
    page_titles: list[str | None] = []

    for page in reader.pages:
        raw = page.extract_text() or ""
        lines = [l.strip() for l in raw.split("\n")]
        title_found = None
        clean_lines: list[str] = []
        for line in lines:
            if not line or _is_header_footer(line):
                continue
            if title_found is None and _is_chapter_title(line):
                title_found = line
            clean_lines.append(line)
        page_texts.append("\n".join(clean_lines))
        page_titles.append(title_found)

    # 2. Group pages into chunks of PAGE_GROUP
    chapters: list[Chapter] = []
    for i in range(0, total_pages, PAGE_GROUP):
        group_pages = page_texts[i: i + PAGE_GROUP]
        group_titles = page_titles[i: i + PAGE_GROUP]
        content = "\n".join(group_pages).strip()
        if not content:
            continue

        # Use first detected chapter title in this group, else page range
        label = next((t for t in group_titles if t), None)
        if label is None:
            label = f"第{i+1}-{min(i+PAGE_GROUP, total_pages)}页"

        chapters.append(Chapter(
            chapter_id=f"ch_{len(chapters):03d}",
            title=label,
            page_start=i + 1,
            page_end=min(i + PAGE_GROUP, total_pages),
            content=content,
            char_count=len(content),
        ))

    total_chars = sum(c.char_count for c in chapters)

    return Textbook(
        textbook_id=tid,
        filename=path.name,
        title=path.stem,
        total_pages=total_pages,
        total_chars=total_chars,
        chapters=chapters,
    )


def _merge_small_chapters(chapters: list[Chapter]) -> list[Chapter]:
    """Merge chapters with too little content into the next chapter."""
    if not chapters:
        return chapters
    merged: list[Chapter] = []
    pending: Chapter | None = None
    for ch in chapters:
        if pending is not None:
            # 把 pending 合并进当前章
            new_content = pending.content + "\n" + ch.content
            ch = ch.model_copy(update={
                "content": new_content,
                "char_count": len(new_content),
                "page_start": pending.page_start,
                "title": ch.title,
            })
            pending = None
        if ch.char_count < MIN_CHAPTER_CHARS:
            pending = ch
        else:
            merged.append(ch)
    if pending is not None:
        merged.append(pending)
    # 重新编号
    for i, ch in enumerate(merged):
        merged[i] = ch.model_copy(update={"chapter_id": f"ch_{i:03d}"})
    return merged


def parse_markdown(file_path: str, textbook_id: Optional[str] = None) -> Textbook:
    path = Path(file_path)
    tid = textbook_id or str(uuid.uuid4())[:8]
    text = path.read_text(encoding="utf-8")

    chapter_re = re.compile(r"^#{1,3}\s+(.+)$", re.MULTILINE)
    matches = list(chapter_re.finditer(text))

    chapters: list[Chapter] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()
        chapters.append(Chapter(
            chapter_id=f"ch_{i:03d}",
            title=m.group(1).strip(),
            page_start=i + 1,
            page_end=i + 1,
            content=content,
            char_count=len(content),
        ))

    if not chapters:
        chapters.append(Chapter(
            chapter_id="ch_000",
            title=path.stem,
            page_start=1,
            page_end=1,
            content=text,
            char_count=len(text),
        ))

    total_chars = sum(c.char_count for c in chapters)
    return Textbook(
        textbook_id=tid,
        filename=path.name,
        title=path.stem,
        total_pages=len(chapters),
        total_chars=total_chars,
        chapters=chapters,
    )


def parse_txt(file_path: str, textbook_id: Optional[str] = None) -> Textbook:
    path = Path(file_path)
    tid = textbook_id or str(uuid.uuid4())[:8]
    text = path.read_text(encoding="utf-8", errors="ignore")

    lines = text.split("\n")
    chapters: list[Chapter] = []
    current_title = path.stem
    current_lines: list[str] = []
    chapter_idx = 0
    line_num = 0

    for line in lines:
        line_num += 1
        if _is_chapter_title(line):
            if current_lines:
                content = "\n".join(current_lines).strip()
                chapters.append(Chapter(
                    chapter_id=f"ch_{chapter_idx:03d}",
                    title=current_title,
                    page_start=max(1, line_num - len(current_lines)),
                    page_end=line_num,
                    content=content,
                    char_count=len(content),
                ))
                chapter_idx += 1
            current_title = line.strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        content = "\n".join(current_lines).strip()
        chapters.append(Chapter(
            chapter_id=f"ch_{chapter_idx:03d}",
            title=current_title,
            page_start=1,
            page_end=line_num,
            content=content,
            char_count=len(content),
        ))

    total_chars = sum(c.char_count for c in chapters)
    return Textbook(
        textbook_id=tid,
        filename=path.name,
        title=path.stem,
        total_pages=1,
        total_chars=total_chars,
        chapters=chapters,
    )


def parse_docx(file_path: str, textbook_id: Optional[str] = None) -> Textbook:
    """Best-effort DOCX parsing: treat each Heading 1/2 as a chapter."""
    from docx import Document  # lazy import; python-docx is in requirements.txt

    path = Path(file_path)
    tid = textbook_id or str(uuid.uuid4())[:8]
    doc = Document(file_path)

    chapters: list[Chapter] = []
    cur_title = path.stem
    cur_lines: list[str] = []

    def _flush(idx: int):
        if not cur_lines:
            return
        content = "\n".join(cur_lines).strip()
        if not content:
            return
        chapters.append(Chapter(
            chapter_id=f"ch_{idx:03d}",
            title=cur_title,
            page_start=idx + 1,
            page_end=idx + 1,
            content=content,
            char_count=len(content),
        ))

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or _is_chapter_title(text):
            _flush(len(chapters))
            cur_title = text
            cur_lines = []
        else:
            cur_lines.append(text)
    _flush(len(chapters))

    if not chapters:
        full = "\n".join((p.text or "").strip() for p in doc.paragraphs if p.text)
        chapters.append(Chapter(
            chapter_id="ch_000",
            title=path.stem,
            page_start=1,
            page_end=1,
            content=full,
            char_count=len(full),
        ))

    total_chars = sum(c.char_count for c in chapters)
    return Textbook(
        textbook_id=tid,
        filename=path.name,
        title=path.stem,
        total_pages=len(chapters),
        total_chars=total_chars,
        chapters=chapters,
    )


PARSERS = {
    ".pdf": parse_pdf,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_txt,
    ".docx": parse_docx,
}


def parse_file(file_path: str, textbook_id: Optional[str] = None) -> Textbook:
    ext = Path(file_path).suffix.lower()
    parser = PARSERS.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file format: {ext}")
    return parser(file_path, textbook_id)
